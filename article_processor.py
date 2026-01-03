try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except ImportError:
    tk = None
    filedialog = None
    messagebox = None
import os
import re
import requests
from bs4 import BeautifulSoup
from markdownify import markdownify as md
from docx import Document
from docx.shared import Pt, Inches
import time
import io
from urllib.parse import urljoin

def parse_txt_file(filepath):
    """
    解析爬虫生成的 txt 文件
    支持两种格式：
    1. 爬虫生成的标准格式 (包含分隔符、标题、链接)
    2. 纯链接列表 (每行一个 URL)
    返回列表: [{'title': ..., 'date': ..., 'link': ...}, ...]
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    items = []
    separator = "-" * 50
    
    if separator in content:
        # 标准格式
        raw_items = content.split(separator)
        for raw in raw_items:
            raw = raw.strip()
            if not raw:
                continue
                
            lines = raw.split('\n')
            item = {'title': '', 'date': '', 'link': ''}
            
            for line in lines:
                line = line.strip()
                if line.startswith("标题:"):
                    item['title'] = line[3:].strip()
                elif line.startswith("链接:"):
                    item['link'] = line[3:].strip()
                elif re.match(r'\d{4}-\d{2}-\d{2}', line): # 简单匹配日期格式
                    item['date'] = line
            
            if item['link']:
                items.append(item)
    else:
        # 纯链接格式
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            # 简单过滤有效链接
            if line and (line.startswith("http://") or line.startswith("https://")):
                items.append({
                    'title': line, # 暂时使用链接作为标题
                    'date': '',
                    'link': line
                })
            
    return items

def fetch_article_content(url):
    """
    抓取 URL 内容，提取正文，转换为 Markdown
    保留图片标记 ![alt](src)，去除普通链接 [text](url)
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        # 降低超时时间，避免卡死
        response = requests.get(url, headers=headers, timeout=10)
        response.encoding = response.apparent_encoding if response.apparent_encoding else 'utf-8'
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 尝试定位正文
        selectors = [
            {'id': re.compile('vsb_content')}, 
            {'class_': 'v_news_content'},      
            {'class_': 'article-2'},           
            {'class_': 'content-box'},         
            {'class_': 'bodytext'},            
            {'class_': 'article-content'},     
            {'class_': 'main-content'},        
            {'name': 'article'},               
        ]
        
        content_div = None
        for selector in selectors:
            if 'name' in selector:
                found = soup.find(selector['name'])
            else:
                found = soup.find('div', **selector)
            
            if found and len(found.get_text().strip()) > 10:
                content_div = found
                break
            
        if not content_div:
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            content_div = soup.body
            
        if not content_div:
            return "未找到正文内容"

        # --- 1. HTML 预处理：移除干扰元素 ---
        # 移除 script, style
        for tag in content_div(["script", "style"]):
            tag.decompose()

        # 移除分享按钮、点击量等干扰信息
        # 根据 debug 结果，分享按钮在 class 为 bshare-custom 或包含 share 的 div/a 中
        # 点击量在 time 标签或包含 "点击" 的文本中
        
        # 移除特定 class 的元素
        for tag in content_div.find_all(class_=re.compile(r'share|fenxiang|bshare', re.I)):
            tag.decompose()
            
        # 移除包含特定关键词的短文本节点 (如 "662次点击", "微信", "QQ空间")
        # 遍历所有文本节点，如果包含关键词且长度较短，则移除其父元素(如果是行内元素)
        for text_node in content_div.find_all(string=True):
            text = text_node.strip()
            if not text:
                continue
                
            # 关键词列表
            keywords = ["次点击", "次浏览", "QQ空间", "新浪微博", "QQ好友", "bshare"]
            
            # 检查是否包含关键词且长度较短，或者是以"上一篇"/"下一篇"开头的导航链接
            is_garbage = (any(k in text for k in keywords) and len(text) < 50)
            is_nav = (text.startswith("上一篇") or text.startswith("下一篇")) and len(text) < 200
            
            if is_garbage or is_nav:
                parent = text_node.parent
                # 如果父元素是 block 元素，可能误删，只删除行内元素或特定标签
                if parent.name in ['span', 'a', 'time', 'em', 'i', 'b', 'strong', 'small']:
                    parent.decompose()
                elif parent.name in ['div', 'p', 'li'] and len(parent.get_text().strip()) < (200 if is_nav else 50):
                    # 如果是块级元素但内容很少（只有垃圾信息），也删除
                    parent.decompose()

        # 处理图片 src，确保是绝对路径
        for img in content_div.find_all('img'):
            src = img.get('src')
            if src:
                img['src'] = urljoin(url, src)
            
        # 转换为 Markdown
        # strip 参数指定要移除格式但保留内容的标签
        # 移除 h1-h6 避免生成 ##, 移除 b/strong 避免生成 **, 移除 a 避免生成链接(保留文字)
        strip_tags = ['script', 'style', 'b', 'strong', 'em', 'i', 'u', 'a', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'font']
        markdown_text = md(str(content_div), heading_style="ATX", strip=strip_tags)
        
        # --- 清理 URL ---
        
        # 1. 处理普通 Markdown 链接 [text](url) -> text
        # 由于上面 strip=['a']，大部分链接已经变成了纯文本，但 markdownify 有时行为不完全一致，保留此正则作为兜底
        markdown_text = re.sub(r'(?<!\!)\[([^\]]+)\]\([^)]+\)', r'\1', markdown_text)
        
        # 2. 再次清理可能残留的 Markdown 符号 (如表格符号 | 或其他)
        # 这里的需求主要是去除 ** 和 ##
        markdown_text = markdown_text.replace("**", "").replace("##", "")
        
        # 3. 清理多余空行
        
        # 2. 移除裸露的 URL (http/https 开头)，但要小心不要误删图片链接中的 URL
        # 图片链接格式: ![...](http...)
        # 我们可以先将图片链接保护起来，或者使用更复杂的正则
        # 简单策略：只匹配前后有空白的 URL，或者不在括号内的 URL
        # 这里简化处理：暂不移除裸露 URL，以免误伤图片地址。
        # 如果必须移除，可以先提取图片，替换为占位符，清理后再还原。
        
        # 3. 清理多余空行
        # 将连续的换行符（包括含有空白字符的）替换为单个段落分隔符
        markdown_text = re.sub(r'\n(\s*\n)+', '\n\n', markdown_text)
        
        return markdown_text.strip()
        
    except Exception as e:
        return f"抓取失败: {e}"

def add_markdown_content_to_doc(doc, markdown_text, progress_callback=None, current_status=None, stop_event=None, download_images=True):
    """
    解析 Markdown 文本，将文字和图片分别添加到 Word 文档
    current_status: (current_index, total_count, title) 用于更新进度
    """
    # 正则匹配图片: ![alt](src)
    # split 会返回 [text, alt, src, text, alt, src, ...]
    parts = re.split(r'!\[([^\]]*)\]\(([^)]+)\)', markdown_text)
    
    # 统计图片总数
    total_images = markdown_text.count('![')
    current_image_idx = 0
    
    # parts[0] 是第一段文字
    # parts[1] 是第一个图片的 alt
    # parts[2] 是第一个图片的 src
    # parts[3] 是第二段文字
    # ...
    
    i = 0
    while i < len(parts):
        # 移除内部的 stop_event 检查，确保当前文章完整处理
        # if stop_event and stop_event.is_set():
        #     return

        text = parts[i].strip()
        if text:
            # 将文本按行分割，每一行作为一个独立的段落添加
            # 使用 splitlines() 可以自动处理 \n, \r, \v (垂直制表符/软回车) 等各种换行符
            for line in text.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        
        # 检查是否有图片
        if i + 2 < len(parts):
            alt = parts[i+1]
            src = parts[i+2]
            current_image_idx += 1
            
            # 更新进度提示
            if progress_callback and current_status:
                idx, total, title = current_status
                if download_images:
                    progress_callback(idx, total, f"正在下载图片 ({current_image_idx}/{total_images}): {title}...")
                else:
                    progress_callback(idx, total, f"正在处理图片链接 ({current_image_idx}/{total_images}): {title}...")

            if not download_images:
                # 仅保存链接
                p = doc.add_paragraph()
                p.add_run(f"[图片: {alt}]").italic = True
                p.add_run(f"({src})").italic = True
            else:
                # 下载并嵌入图片
                try:
                    # print(f"正在下载图片: {src}")
                    headers = {
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                    }
                    # 降低图片下载超时时间，使用 stream=True 以便检查大小
                    # timeout=(connect, read)
                    with requests.get(src, headers=headers, timeout=(3, 5), stream=True) as img_response:
                        if img_response.status_code == 200:
                            # 检查 Content-Length (限制 10MB)
                            content_length = img_response.headers.get('content-length')
                            if content_length and int(content_length) > 10 * 1024 * 1024:
                                p = doc.add_paragraph()
                                p.add_run(f"[图片过大 ({int(content_length)/1024/1024:.1f}MB)，已跳过]: {src}").italic = True
                            else:
                                # 读取内容
                                image_content = img_response.content
                                image_stream = io.BytesIO(image_content)
                                # 插入图片，限制宽度，避免溢出
                                doc.add_picture(image_stream, width=Inches(5.5))
                        else:
                            p = doc.add_paragraph()
                            p.add_run(f"[图片下载失败 ({img_response.status_code}): {src}]").italic = True
                except Exception as e:
                    # 图片下载或插入失败时，保留链接
                    p = doc.add_paragraph()
                    p.add_run(f"[图片插入错误: {e}]").italic = True
                    p.add_run(f" 链接: {src}")
            
            i += 3 # 跳过 alt 和 src
        else:
            i += 1

def generate_word_doc(items, output_path, max_size_mb=100, progress_callback=None, stop_event=None, pause_event=None, start_index=1, download_images=True):
    doc = Document()
    
    # 设置默认字体 (可选)
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)
    
    total = len(items)
    current_part = 1
    base_name, ext = os.path.splitext(output_path)
    has_content = False
    
    # 调整切片，start_index 是 1-based
    process_items = items[start_index-1:]
    
    for i, item in enumerate(process_items):
        real_index = start_index + i
        
        # 检查停止信号
        if stop_event and stop_event.is_set():
            if progress_callback:
                progress_callback(real_index - 1, total, "任务已终止")
            break
            
        # 检查暂停信号
        if pause_event:
            while pause_event.is_set():
                if stop_event and stop_event.is_set():
                    break
                time.sleep(0.5)
        
        if progress_callback:
            progress_callback(real_index, total, item['title'])
            
        # 添加标题
        doc.add_heading(item['title'], level=1)
        
        # 添加元数据
        p = doc.add_paragraph()
        p.add_run(f"发布日期: {item['date']}\n").bold = True
        p.add_run(f"原文链接: {item['link']}") # 这里保留原文链接
        
        # 抓取内容
        content = fetch_article_content(item['link'])
        
        # 添加正文
        doc.add_heading('文章内容:', level=2)
        # 使用新的处理函数
        add_markdown_content_to_doc(doc, content, progress_callback, (real_index, total, item['title']), stop_event, download_images)
        
        # 分页
        doc.add_page_break()
        has_content = True
        
        # 检查文件大小
        try:
            # 保存到内存流以检查大小
            buffer = io.BytesIO()
            doc.save(buffer)
            size_bytes = buffer.tell()
            
            if size_bytes > max_size_mb * 1024 * 1024:
                # 超过大小，保存当前部分
                part_path = f"{base_name}_part{current_part}{ext}"
                with open(part_path, "wb") as f:
                    f.write(buffer.getvalue())
                
                if progress_callback:
                    progress_callback(real_index, total, f"已保存分卷: {os.path.basename(part_path)}")
                
                # 重置文档
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = '微软雅黑'
                font.size = Pt(10)
                
                current_part += 1
                has_content = False
        except Exception as e:
            print(f"Size check error: {e}")
        
        # 避免请求过快
        time.sleep(0.5)
        
    # 保存剩余内容
    if has_content:
        if current_part == 1:
            doc.save(output_path)
        else:
            part_path = f"{base_name}_part{current_part}{ext}"
            doc.save(part_path)

class ProcessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("文章提取与Word生成工具")
        self.root.geometry("500x400")
        
        self.create_widgets()
        
    def create_widgets(self):
        frame = tk.Frame(self.root, padx=20, pady=20)
        frame.pack(fill="both", expand=True)
        
        tk.Label(frame, text="1. 选择爬虫生成的 txt 文件:").pack(anchor="w")
        
        self.file_path_var = tk.StringVar()
        entry_frame = tk.Frame(frame)
        entry_frame.pack(fill="x", pady=5)
        tk.Entry(entry_frame, textvariable=self.file_path_var).pack(side="left", fill="x", expand=True)
        tk.Button(entry_frame, text="浏览", command=self.browse_file).pack(side="right", padx=5)
        
        tk.Label(frame, text="2. 设置单个文档最大大小 (MB):").pack(anchor="w", pady=(10, 0))
        self.size_var = tk.StringVar(value="100")
        tk.Entry(frame, textvariable=self.size_var).pack(fill="x", pady=5)
        
        tk.Label(frame, text="3. 生成 Word 文档:").pack(anchor="w", pady=(15, 0))
        tk.Button(frame, text="开始处理并保存", command=self.start_process, bg="#dddddd").pack(fill="x", pady=5)
        
        self.log_text = tk.Text(frame, height=8, state="disabled")
        self.log_text.pack(fill="both", expand=True, pady=10)
        
    def log(self, msg):
        self.log_text.config(state="normal")
        self.log_text.insert("end", msg + "\n")
        self.log_text.see("end")
        self.log_text.config(state="disabled")
        self.root.update()

    def browse_file(self):
        f = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if f:
            self.file_path_var.set(f)
            
    def start_process(self):
        input_path = self.file_path_var.get()
        if not input_path or not os.path.exists(input_path):
            messagebox.showerror("错误", "请选择有效的 txt 文件")
            return
            
        try:
            max_size = float(self.size_var.get())
        except ValueError:
            messagebox.showerror("错误", "请输入有效的大小数值")
            return
            
        # 解析
        self.log("正在解析文件...")
        items = parse_txt_file(input_path)
        self.log(f"找到 {len(items)} 篇文章链接")
        
        if not items:
            self.log("未找到有效链接，请检查文件格式")
            return
            
        # 选择保存路径
        output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not output_path:
            return
            
        self.log(f"开始抓取并生成 Word，保存至: {output_path}")
        self.log(f"单个文档最大限制: {max_size} MB")
        
        try:
            generate_word_doc(items, output_path, max_size, self.progress_update)
            self.log("处理完成！")
            messagebox.showinfo("完成", f"文档已保存")
        except Exception as e:
            self.log(f"发生错误: {e}")
            messagebox.showerror("错误", str(e))
            
    def progress_update(self, current, total, title):
        self.log(f"[{current}/{total}] 处理: {title}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ProcessorApp(root)
    root.mainloop()
