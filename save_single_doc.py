import sys
import os
from docx import Document
from docx.shared import Pt
import time

# Add current directory to path
sys.path.append(os.getcwd())
from article_processor import fetch_article_content, add_markdown_content_to_doc

url = "https://www.sdxd.edu.cn/detail/202512121125523yri14xwaz2ch4soy4.html"
title = "测试带图片文章" 
date = "2025-12-12"

print(f"Generating Word doc for: {url}")

content = fetch_article_content(url)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)

doc.add_heading(title, level=1)
p = doc.add_paragraph()
p.add_run(f"发布日期: {date}\n").bold = True
p.add_run(f"原文链接: {url}")

doc.add_heading('文章内容:', level=2)
# Use the new function to handle images
add_markdown_content_to_doc(doc, content)

output_path = "d:\\pachong\\single_article_with_links.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
