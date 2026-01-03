import os
import json
import requests
import time
from docx import Document
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ================= 配置区域 =================
# 请在此处填写您的 DeepSeek API Key
API_KEY = "YOUR_DEEPSEEK_API_KEY_HERE" 

# DeepSeek API 地址 (通常兼容 OpenAI 格式)
API_URL = "https://api.deepseek.com/chat/completions"

# 模型名称 (deepseek-chat 或 deepseek-reasoner)
MODEL_NAME = "deepseek-chat"

# 输出文件前缀
OUTPUT_FILE_PREFIX = "generated_qa_dataset"

# 最大分片长度 (字符数)
# 降低到 15000 以避免超时
MAX_CHUNK_SIZE = 15000 

# 单个 Word 文件最大大小 (字节) - 100MB
MAX_FILE_SIZE_BYTES = 100 * 1024 * 1024 
# ===========================================

# 提示词 (System Prompt)
SYSTEM_PROMPT = """你是一个专业的知识库构建助手。你的任务是根据用户提供的文档内容，提取并生成高质量的问答对（QA pairs）。
这些问答对将用于训练智能客服或知识库大模型。

请遵循以下原则：
1.准确性：问题的答案必须完全基于提供的文档内容，不得编造。
2.覆盖面：尽可能覆盖文档中的关键信息点、概念、流程和规则。
3.独立性：每个问题和答案应当是独立的，不需要上下文也能理解。
4.格式：输出必须是标准的 JSON 格式，包含一个列表，每个元素是一个对象，包含 "question" 和 "answer" 两个字段。

示例输出格式：
[
    {
        "question": "什么是...",
        "answer": "..."
    },
    {
        "question": "如何办理...",
        "answer": "..."
    }
]
"""

class OutputManager:
    def __init__(self, prefix, max_size):
        self.prefix = prefix
        self.max_size = max_size
        self.file_index = 1
        self.current_doc = Document()
        self.current_qa_count = 0
        self.current_file_path = self._get_file_path()

    def _get_file_path(self):
        return f"{self.prefix}_{self.file_index}.docx"

    def add_qa_list(self, qa_list, source_file):
        if not qa_list:
            return

        # 如果是新文件或者刚切换了源文件，添加一次来源标题
        if not hasattr(self, 'last_source_file') or self.last_source_file != source_file:
            self.current_doc.add_heading(f"来源文档: {source_file}", level=2)
            self.last_source_file = source_file

        for item in qa_list:
            # 去除首尾空白字符和换行符
            q = item.get('question', '').strip()
            a = item.get('answer', '').strip()
            
            if not q or not a:
                continue

            # 添加到 Word 文档
            self.current_doc.add_paragraph(f"Q: {q}", style='List Number')
            self.current_doc.add_paragraph(f"A: {a}")
            # 增加一个空行作为分隔，替代原来的横线
            self.current_doc.add_paragraph("")
            
            self.current_qa_count += 1

        # 保存当前进度
        self.save()
        
        # 检查文件大小是否超过限制
        if os.path.exists(self.current_file_path):
            file_size = os.path.getsize(self.current_file_path)
            if file_size > self.max_size:
                print(f"当前文件 {self.current_file_path} 大小 ({file_size/1024/1024:.2f} MB) 超过限制，创建新文件...")
                self.file_index += 1
                self.current_doc = Document()
                self.current_qa_count = 0
                self.current_file_path = self._get_file_path()

    def save(self):
        try:
            self.current_doc.save(self.current_file_path)
            # print(f"已保存进度到: {self.current_file_path}")
        except Exception as e:
            print(f"保存文件失败: {e}")

def read_word_file(file_path):
    """读取 Word 文档内容"""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            # 去除段落前后的空白字符
            text = para.text.strip()
            # 过滤掉空段落
            if text:
                full_text.append(text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"读取文件 {file_path} 失败: {e}")
        return None

def split_text(text, max_length=MAX_CHUNK_SIZE):
    """将长文本切分为多个片段，防止超过 Token 限制"""
    chunks = []
    current_chunk = []
    current_length = 0
    
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 如果单段落本身就超长（极少见），强制切分
        if len(para) > max_length:
            # 先把当前累积的存入
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_length = 0
            # 强制切分长段落
            for i in range(0, len(para), max_length):
                chunks.append(para[i:i+max_length])
            continue

        if current_length + len(para) > max_length:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_length = len(para)
        else:
            current_chunk.append(para)
            current_length += len(para) + 1 # +1 for newline
            
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
        
    return chunks

def create_session_with_retries():
    """创建带有重试机制的 requests session"""
    session = requests.Session()
    retries = Retry(
        total=3,
        backoff_factor=1,
        status_forcelist=[500, 502, 503, 504],
        allowed_methods=["POST"]
    )
    adapter = HTTPAdapter(max_retries=retries)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    return session

def generate_qa(text, session):
    """调用 DeepSeek API 生成 QA"""
    if not text.strip():
        return None

    if not API_KEY:
        print("错误: 未设置 API_KEY。请在脚本中填写您的 DeepSeek API Key。")
        return None

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }

    # 构造 Prompt
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": f"请根据以下文档内容生成 QA 对：\n\n{text}"}
    ]

    data = {
        "model": MODEL_NAME,
        "messages": messages,
        "temperature": 0.7, # 控制生成的多样性
        "stream": False,
        "response_format": { "type": "json_object" } # 强制 JSON 输出 (如果模型支持)
    }

    try:
        # print("正在调用 DeepSeek API...")
        # 增加超时时间到 180 秒
        response = session.post(API_URL, headers=headers, json=data, timeout=180)
        response.raise_for_status()
        
        result = response.json()
        if 'choices' in result and len(result['choices']) > 0:
            content = result['choices'][0]['message']['content']
            return content
        else:
            print(f"API 返回格式异常: {result}")
            return None
            
    except requests.exceptions.RequestException as e:
        print(f"API 请求失败: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"响应状态码: {e.response.status_code}")
            print(f"响应内容: {e.response.text}")
        return None
    except Exception as e:
        print(f"发生未知错误: {e}")
        return None

def parse_json_response(content):
    """解析 API 返回的 JSON 字符串"""
    try:
        # 清理可能存在的 Markdown 代码块标记
        clean_content = content.replace("```json", "").replace("```", "").strip()
        return json.loads(clean_content)
    except json.JSONDecodeError:
        print("解析 JSON 失败，尝试修复或返回原始文本...")
        # 这里可以添加更复杂的修复逻辑，或者直接返回 None
        return None

def main():
    print("=== Word 文档转 QA 知识库脚本 (Word 输出版) ===")
    
    input_files = []
    
    # 交互式输入文件路径
    print("请输入 Word 文档路径 (支持拖入文件，输入 'done' 或直接回车结束输入):")
    while True:
        user_input = input("文件路径 > ").strip()
        if user_input.lower() == 'done' or user_input == '':
            if input_files: # 如果已经有文件了，回车表示结束
                break
            elif user_input == '': # 如果还没文件就回车，继续提示
                continue
        
        # 去除可能存在的引号（Windows 拖入文件时可能会带引号）
        file_path = user_input.strip('"').strip("'")
        
        if os.path.exists(file_path):
            if file_path.endswith('.docx'):
                if file_path not in input_files:
                    input_files.append(file_path)
                    print(f"已添加: {file_path}")
                else:
                    print("文件已在列表中。")
            else:
                print("目前仅支持 .docx 格式的文件。")
        else:
            print("文件不存在，请检查路径。")

    if not input_files:
        print("未选择任何文件，程序退出。")
        return

    # 初始化输出管理器
    output_manager = OutputManager(OUTPUT_FILE_PREFIX, MAX_FILE_SIZE_BYTES)
    session = create_session_with_retries()

    total_qa_count = 0

    for file_idx, file_path in enumerate(input_files):
        print(f"\n[{file_idx+1}/{len(input_files)}] 正在处理文件: {os.path.basename(file_path)} ...")
        
        # 1. 读取内容
        text = read_word_file(file_path)
        if not text:
            continue
            
        print(f"文档读取成功，长度: {len(text)} 字符")
        
        # 2. 切分文本并生成 QA
        chunks = split_text(text)
        total_chunks = len(chunks)
        if total_chunks > 1:
            print(f"文档较长，已自动切分为 {total_chunks} 个片段进行处理。")
        
        for i, chunk in enumerate(chunks):
            print(f"  -> 正在处理片段 {i+1}/{total_chunks} (长度: {len(chunk)})... ", end="", flush=True)
            
            start_time = time.time()
            qa_content = generate_qa(chunk, session)
            elapsed_time = time.time() - start_time
            
            if qa_content:
                # 3. 解析结果
                qa_json = parse_json_response(qa_content)
                
                if qa_json:
                    # 确保是列表格式
                    if isinstance(qa_json, dict):
                        for key in qa_json:
                            if isinstance(qa_json[key], list):
                                qa_json = qa_json[key]
                                break
                        else:
                            qa_json = [qa_json]
                    
                    if isinstance(qa_json, list):
                        count = len(qa_json)
                        # 4. 写入 Word
                        output_manager.add_qa_list(qa_json, os.path.basename(file_path))
                        total_qa_count += count
                        print(f"成功! 耗时 {elapsed_time:.1f}s, 提取 {count} 个 QA 对。")
                    else:
                        print("失败 (格式错误)。")
                else:
                    print("失败 (解析错误)。")
            else:
                print("失败 (API 错误)。")

    print(f"\n===========================================")
    print(f"全部完成！共生成 {total_qa_count} 个 QA 对。")
    print(f"结果已保存至: {os.path.abspath(output_manager.current_file_path)} (及之前的分卷)")
    print(f"===========================================")

if __name__ == "__main__":
    main()
